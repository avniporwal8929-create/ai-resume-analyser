import pdfplumber
import docx
import io

def extract_text_from_pdf(file_file) -> str:
    """
    Extracts text from a PDF file using pdfplumber.
    Supports file-like objects (e.g. Streamlit's UploadedFile).
    """
    try:
        text = ""
        with pdfplumber.open(file_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if not text.strip():
            raise ValueError("The PDF file contains no selectable text. It might be empty or a scanned image.")
        return text
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def extract_text_from_docx(file_file) -> str:
    """
    Extracts text from a Word Document (.docx) using python-docx.
    Iterates over paragraphs and tables to compile all text.
    """
    try:
        doc = docx.Document(file_file)
        text_lines = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)
                
        # Extract text from tables (often used for resume layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_lines.append(cell_text)
                        
        extracted_text = "\n".join(text_lines)
        if not extracted_text.strip():
            raise ValueError("The DOCX file seems to be empty.")
        return extracted_text
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def extract_text_from_txt(file_file) -> str:
    """
    Extracts text from a plain text (.txt) file.
    Attempts UTF-8 decoding, falling back to Latin-1.
    """
    try:
        file_bytes = file_file.read()
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
    except Exception as e:
        raise ValueError(f"Error parsing TXT file: {str(e)}")

def parse_resume(uploaded_file) -> str:
    """
    Determines file type based on extension and extracts text.
    Resets the file pointer to the beginning before parsing.
    """
    if uploaded_file is None:
        raise ValueError("No file uploaded.")
        
    filename = uploaded_file.name.lower()
    
    # Reset the buffer position
    uploaded_file.seek(0)
    
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    elif filename.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
